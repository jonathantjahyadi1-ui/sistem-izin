import io
import os
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


TEMPLATE_PATH = (
    Path(__file__).resolve().parent
    / 'docx_templates'
    / 'Form_Pengajuan_Pengunduran_Diri.docx'
)

MONTH_NAMES = (
    'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
    'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
)

REASON_LABELS = {
    'pekerjaan_lain': 'Mendapat pekerjaan lain',
    'pendidikan': 'Melanjutkan pendidikan',
    'kesehatan': 'Kesehatan (pribadi/keluarga)',
    'relokasi': 'Relokasi domisili',
    'keluarga': 'Alasan pribadi/keluarga lainnya',
    'lainnya': 'Lainnya',
}


def format_date_id(value):
    if not value:
        return '-'
    return f'{value.day:02d} {MONTH_NAMES[value.month - 1]} {value.year}'


def _set_paragraph_text(paragraph, text, *, bold=None, italic=None):
    """Ganti teks sambil mempertahankan format dasar paragraf template."""
    if paragraph.runs:
        first_run = paragraph.runs[0]
        template_rpr = deepcopy(first_run._element.rPr)
        for run in list(paragraph.runs):
            paragraph._element.remove(run._element)
        run = paragraph.add_run(str(text))
        if template_rpr is not None:
            run._element.insert(0, template_rpr)
    else:
        run = paragraph.add_run(str(text))

    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _set_cell_value(cell, value):
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, value or '-')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _insert_paragraph_before(paragraph, text):
    new_paragraph = paragraph.insert_paragraph_before()
    _set_paragraph_text(new_paragraph, text)
    return new_paragraph


def _add_draft_mark(document):
    # Gunakan baris kosong bawaan template agar label draft tidak mengubah
    # tinggi header atau menutupi konten pada halaman kedua.
    paragraph = document.paragraphs[3]
    _set_paragraph_text(paragraph, 'DRAFT - BELUM DISETUJUI', bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(185, 28, 28)
    run.font.name = 'Arial'


def build_resignation_document(resignation, get_user, *, final=False):
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError('Template Word pengunduran diri tidak ditemukan.')

    document = Document(str(TEMPLATE_PATH))
    company_name = os.getenv('COMPANY_NAME', 'PT Thriva Grovia Bersama')

    if not final:
        _add_draft_mark(document)

    paragraphs = document.paragraphs
    _set_paragraph_text(paragraphs[0], company_name, bold=True)
    _set_paragraph_text(
        paragraphs[6],
        f'Bapak/Ibu {resignation.supervisor_name or "-"} '
        '(Atasan Langsung)'
    )
    _set_paragraph_text(paragraphs[7], company_name)

    employee_table = document.tables[1]
    employee_values = (
        resignation.employee_name,
        resignation.employee_nik,
        resignation.position,
        resignation.employee_division,
        resignation.supervisor_name,
        (
            '☒ Tetap (PKWTT)     ☐ Kontrak (PKWT)'
            if resignation.employment_status == 'pkwtt'
            else '☐ Tetap (PKWTT)     ☒ Kontrak (PKWT)'
            if resignation.employment_status == 'pkwt'
            else '☐ Tetap (PKWTT)     ☐ Kontrak (PKWT)'
        ),
        format_date_id(resignation.start_date),
    )
    for row, value in zip(employee_table.rows, employee_values):
        _set_cell_value(row.cells[1], value)

    resignation_table = document.tables[2]
    notice_text = f'{resignation.notice_days or 0} hari'
    if (resignation.notice_days or 0) < 30:
        notice_text += ' (pengajuan notice kurang dari 30 hari)'
    resignation_values = (
        format_date_id(resignation.submission_date),
        format_date_id(resignation.effective_date),
        notice_text,
    )
    for row, value in zip(resignation_table.rows, resignation_values):
        _set_cell_value(row.cells[1], value)

    if (resignation.notice_days or 0) < 30 and resignation.short_notice_reason:
        section_c_heading = next(
            p for p in document.paragraphs
            if p.text.strip().startswith('C. Alasan')
        )
        note_paragraph = _insert_paragraph_before(
            section_c_heading,
            'Catatan notice kurang dari 30 hari: '
            f'{resignation.short_notice_reason}'
        )
        if note_paragraph.runs:
            note_paragraph.runs[0].italic = True

    selected_reasons = set(resignation.reason_codes)
    reason_paragraphs = paragraphs[14:20]
    reason_keys = (
        'pekerjaan_lain', 'pendidikan', 'kesehatan',
        'relokasi', 'keluarga', 'lainnya'
    )
    for paragraph, reason_key in zip(reason_paragraphs, reason_keys):
        marker = '☒' if reason_key in selected_reasons else '☐'
        label = REASON_LABELS[reason_key]
        if reason_key == 'lainnya':
            label = f'Lainnya: {resignation.reason_other or "-"}'
        _set_paragraph_text(paragraph, f'{marker} {label}')

    # Tambahkan syarat ikatan dinas yang terdapat dalam PP 35/2021.
    section_e_heading = next(
        p for p in document.paragraphs
        if p.text.strip().startswith('E. Informasi')
    )
    _insert_paragraph_before(
        section_e_heading,
        '5. Saya menyatakan tidak sedang terikat ikatan dinas atau telah '
        'menyelesaikan kewajiban ikatan dinas sesuai ketentuan perusahaan.'
    )

    admin_table = document.tables[3]
    admin_values = (
        resignation.correspondence_address,
        resignation.phone_number,
        resignation.personal_email,
        (
            f'{resignation.bank_name or "-"} / '
            f'{resignation.bank_account_number or "-"} / '
            f'a.n. {resignation.bank_account_holder or "-"}'
        ),
    )
    for row, value in zip(admin_table.rows, admin_values):
        _set_cell_value(row.cells[1], value)

    approval_table = document.tables[4]
    employee_cell, supervisor_cell, hrd_cell = approval_table.rows[0].cells

    _set_paragraph_text(employee_cell.paragraphs[2], resignation.employee_name or '-')
    _set_paragraph_text(
        employee_cell.paragraphs[4],
        'Dikonfirmasi melalui sistem',
        italic=True,
    )
    _set_paragraph_text(
        employee_cell.paragraphs[6],
        format_date_id(resignation.submission_date),
    )

    supervisor_user = (
        get_user(resignation.supervisor_by)
        if resignation.supervisor_by else None
    )
    supervisor_approver_name = (
        (supervisor_user.nama_lengkap or supervisor_user.username)
        if supervisor_user else resignation.supervisor_name or '-'
    )
    _set_paragraph_text(supervisor_cell.paragraphs[2], supervisor_approver_name)
    supervisor_approved = resignation.supervisor_decision == 'approved'
    supervisor_rejected = resignation.supervisor_decision == 'rejected'
    _set_paragraph_text(
        supervisor_cell.paragraphs[4],
        (
            f'{"☒" if supervisor_approved else "☐"} Disetujui     '
            f'{"☒" if supervisor_rejected else "☐"} Tidak Disetujui'
        ),
    )
    _set_paragraph_text(
        supervisor_cell.paragraphs[6],
        'Persetujuan melalui sistem' if resignation.supervisor_at else '-',
        italic=True,
    )
    _set_paragraph_text(
        supervisor_cell.paragraphs[8],
        format_date_id(
            resignation.supervisor_at.date()
            if resignation.supervisor_at else None
        ),
    )

    hrd_user = get_user(resignation.hrd_by) if resignation.hrd_by else None
    hrd_name = (
        (hrd_user.nama_lengkap or hrd_user.username) if hrd_user else '-'
    )
    _set_paragraph_text(hrd_cell.paragraphs[2], hrd_name)
    _set_paragraph_text(
        hrd_cell.paragraphs[4],
        format_date_id(
            resignation.hrd_at.date() if resignation.hrd_at else None
        ),
    )
    _set_paragraph_text(
        hrd_cell.paragraphs[6],
        'Divalidasi melalui sistem' if resignation.hrd_at else '-',
        italic=True,
    )
    _set_paragraph_text(
        hrd_cell.paragraphs[8],
        format_date_id(
            resignation.hrd_at.date() if resignation.hrd_at else None
        ),
    )

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def make_document_filename(resignation, *, final=False):
    safe_name = ''.join(
        char if char.isalnum() else '_'
        for char in (resignation.employee_name or 'Karyawan')
    ).strip('_')
    prefix = 'FINAL' if final else 'DRAFT'
    reference = resignation.reference_no or f'RESIGN-{resignation.id}'
    return f'{prefix}_{reference}_{safe_name}.docx'
